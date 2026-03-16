from flask import Flask, request, jsonify, send_file
from docx import Document
import json, io, os, re

app = Flask(__name__)

@app.route('/preventivo', methods=['POST'])
def genera_preventivo():
    raw = request.get_data(as_text=True)
    clean = re.sub(r'```json|```', '', raw).strip()
    clean = clean.replace("\n", " ")

    dati = json.loads(clean)

    doc = Document()
    doc.add_heading('PREVENTIVO', 0)
    doc.add_paragraph(f"Indirizzo: {dati.get('indirizzo','')}")
    doc.add_paragraph(f"Data: {dati.get('data','')}")
    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = 'Descrizione'
    hdr[1].text = 'Persone'
    hdr[2].text = 'Ore'
    hdr[3].text = 'Prezzo/h'
    hdr[4].text = 'Totale'

    for riga in dati.get('righe', []):
        row = table.add_row().cells
        row[0].text = riga.get('descrizione','')
        row[1].text = str(riga.get('persone',1))
        row[2].text = str(riga.get('ore',0))
        row[3].text = str(riga.get('prezzo_ora',0))
        row[4].text = str(riga.get('totale',0))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(buf,
                     as_attachment=True,
                     download_name="preventivo.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route('/')
def home():
    return 'Server preventivi attivo!'

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
